"""Extract plain text from resume files (PDF, DOCX, TXT)."""
from pathlib import Path

SUPPORTED = {".pdf", ".docx", ".txt"}


def extract_text(path: str) -> str:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"File not found: {p}")
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix == ".docx":
        return _extract_docx(p)
    if suffix == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(
        f"Unsupported file type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED))}"
    )


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
